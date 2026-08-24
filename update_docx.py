import docx

def create_docx_from_txt(txt_path, docx_path):
    doc = docx.Document()
    with open(txt_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    for line in lines:
        line = line.strip('\n')
        if line.startswith('----'):
            continue  # Skip raw text dividers
        if line.isupper() and len(line) > 3 and not line.startswith('*'):
            # Heading
            p = doc.add_heading(line, level=1)
        else:
            # Normal text
            doc.add_paragraph(line)

    doc.save(docx_path)
    print("Done")

if __name__ == '__main__':
    create_docx_from_txt('Resume_Project_Details.txt', 'Resume_Project_Details.docx')
